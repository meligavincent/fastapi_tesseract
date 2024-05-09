from typing import Annotated

from fastapi import FastAPI, File, UploadFile
import cv2
import pytesseract
import numpy as np
import docx
import fitz
from pdf2image import convert_from_path 
from docx.enum.section import WD_SECTION

from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles

import os
import shutil
from model import FileProcessing
from pymongo import MongoClient


app = FastAPI()
app.mount("/files", StaticFiles(directory="files"), name="files")
origins = [
    "http://localhost.tiangolo.com",
    "https://localhost.tiangolo.com",
    "http://localhost",
    "http://localhost:8080",
    "http://localhost:3000",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.post("/files/")
async def create_file(file: Annotated[bytes, File()]):
    return {"file_size": len(file)}



def create_upload_files(files: list[UploadFile]):
    for file in files:
        with open(f"files/{file.filename}", "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    return {"filenames": [file.filename for file in files]}

@app.post("/uploadfiles/")
async def create_upload_file(files: list[UploadFile]) -> list[FileProcessing]:

    file_processing = []
    files = create_upload_files(files=files)["filenames"]

    for file in files:
        file = f'./files/{file}'
        if str(file).split(".")[-1] in ["png", "jpg", "PNG", "JPG"]:
    
            text = ocr(file)

        elif str(file).split(".")[-1] in ["pdf",'PDF']:
            read_scanned_pdf(file)
            text = None
        else:
            text = None
        
        try :
            os.mkdir('files')
        except OSError:
            print('files folder already created')


        file_processing.append(
            FileProcessing(
                file_path=f"http://localhost:8000/files/{file}",
                file_name=file,
                processed=True,
                extracted_text=text,
            )
        )

        file_data = FileProcessing(
            file_path=f"http://localhost:8000/files/{file}",
            processed=True,
            extracted_text=text,
        ).dict()
        client = MongoClient("mongodb://localhost:27017/")
        db = client["OCR"]
        collection = db["FileProcessing"]
        collection.insert_one(file_data)
    return file_processing


@app.get("/getfiles/")
async def get_files():
    client = MongoClient("mongodb://localhost:27017/")
    db = client["OCR"]
    collection = db["FileProcessing"]
    data = list(
        collection.find(
            {},
            {
                "_id": 0,
                "file_path": 1,
                "file_name": 1,
                "processed": 1,
                "extracted_text": 1,
            },
        )
    )

    return data


def valid_xml_char_ordinal(c):

    codepoint = ord(c)
    # conditions ordered by presumed frequency
    return (
        0x20 <= codepoint <= 0xD7FF or
        codepoint in (0x9, 0xA, 0xD) or
        0xE000 <= codepoint <= 0xFFFD or
        0x10000 <= codepoint <= 0x10FFFF
        )

def ocr(file):
    img = cv2.imread(file)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    threshold_img = cv2.threshold(
        gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU
    )[1]
    text = pytesseract.image_to_string(threshold_img)
    text = ''.join(c for c in text if valid_xml_char_ordinal(c))
    write_word_file(text)
    return text

def read_pdf(pdf_file):
    pdf_document = fitz.open(pdf_file)

    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        text = page.get_text()
        
        print(f"Page {page_num + 1}: {text}")

    pdf_document.close()

    return None

def read_scanned_pdf(pdf_file):

    pages = convert_from_path(pdf_file, 500,poppler_path='C:\\Users\\GENERAL-STORES\\Downloads\\poppler-23.08.0\\Library\\bin')
    doc = docx.Document()

    for count, page in enumerate(pages):
        page.save(f'out{count}.jpg', 'JPEG')
        # Add a new page break
        
        # Add a new paragraph on the new page
        new_paragraph = doc.add_paragraph()
        # Write text to the new paragraph
        new_paragraph.text = ocr(f'out{count}.jpg')
        doc.add_page_break()
        os.remove(f'out{count}.jpg')

    doc.save("./files/word_output/new_document.docx")
    print("finished")
    return doc


def write_word_file(text):
    doc = docx.Document()

    doc.add_paragraph(text=text)
    doc.save('demo.docx')
    return None